import streamlit as st
import json
import os
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def load_css():
    """Load custom CSS styles for the application"""
    css = """
    <style>
    /* Main theme variables */
    :root {
        --primary-color: #7c3aed;
        --primary-gradient: linear-gradient(135deg, #7c3aed 0%, #6d28d9 100%);
        --secondary-color: #64748b;
        --accent-color: #f59e0b;
        --success-color: #10b981;
        --warning-color: #f59e0b;
        --error-color: #ef4444;

        --bg-primary: #ffffff;
        --bg-secondary: #f8fafc;
        --bg-tertiary: #f1f5f9;

        --text: #1e293b;
        --text-secondary: #64748b;
        --text-muted: #94a3b8;

        --glass-bg: rgba(255, 255, 255, 0.95);
        --glass-border: rgba(255, 255, 255, 0.2);

        --shadow-light: 0 1px 3px rgba(0, 0, 0, 0.1);
        --shadow-medium: 0 4px 6px rgba(0, 0, 0, 0.1);
        --shadow-heavy: 0 10px 15px rgba(0, 0, 0, 0.1);

        --border-radius: 12px;
        --border-radius-small: 8px;
        --border-radius-large: 16px;

        --transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        --transition-fast: all 0.15s cubic-bezier(0.4, 0, 0.2, 1);
        --transition-slow: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);

        --spacing-xs: 4px;
        --spacing-sm: 8px;
        --spacing-md: 16px;
        --spacing-lg: 24px;
        --spacing-xl: 32px;
        --spacing-2xl: 48px;
    }

    /* Dark theme variables */
    [data-theme="dark"] {
        --bg-primary: #0f172a;
        --bg-secondary: #1e293b;
        --bg-tertiary: #334155;

        --text: #f8fafc;
        --text-secondary: #cbd5e1;
        --text-muted: #64748b;

        --glass-bg: rgba(15, 23, 42, 0.95);
        --glass-border: rgba(255, 255, 255, 0.1);

        --shadow-light: 0 1px 3px rgba(0, 0, 0, 0.3);
        --shadow-medium: 0 4px 6px rgba(0, 0, 0, 0.3);
        --shadow-heavy: 0 10px 15px rgba(0, 0, 0, 0.3);
    }

    /* Global styles */
    * {
        box-sizing: border-box;
    }

    body {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        background: var(--bg-primary);
        color: var(--text);
        line-height: 1.6;
        margin: 0;
        padding: 0;
        transition: var(--transition);
    }

    /* Streamlit overrides */
    .stApp {
        margin: 0 !important;
        padding: 0 !important;
    }

    .stMain {
        padding: 0 !important;
    }

    .stSidebar {
        padding: 0 !important;
        margin: 0 !important;
        background: var(--bg-primary) !important;
    }

    .stSidebarContent {
        padding: 0 !important;
    }

    /* Main header */
    .main-header {
        text-align: center;
        margin: 0;
        padding: var(--spacing-lg) var(--spacing-xl);
        background: var(--primary-gradient);
        border-radius: var(--border-radius) var(--border-radius) 0 0;
    }

    .main-title {
        font-size: 2.5rem;
        font-weight: 700;
        color: white;
        margin: 0;
        line-height: 1;
    }

    .subtitle {
        font-size: 1.125rem;
        color: white;
        margin: 0;
        line-height: 1;
        font-weight: 400;
    }

    /* Sidebar styles */
    .sidebar-header {
        margin-bottom: var(--spacing-lg);
        padding: var(--spacing-md);
        background: var(--primary-gradient);
        border-radius: var(--border-radius);
        margin-bottom: var(--spacing-lg);
    }

    .sidebar-header h3 {
        color: white;
        font-size: 1.25rem;
        font-weight: 600;
        margin: 0;
        text-align: center;
    }

    /* Chat container */
    .chat-container {
        min-height: 400px;
        padding: 0;
        margin: -20px 0 0 0;
        background: transparent;
        border: none;
    }

    /* Chat messages */
    .chat-message {
        display: flex;
        margin-bottom: var(--spacing-lg);
        padding: var(--spacing-md);
        border-radius: var(--border-radius);
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        border: 1px solid var(--glass-border);
        box-shadow: var(--shadow-light);
        transition: var(--transition);
    }

    .chat-message:hover {
        transform: translateY(-1px);
        box-shadow: var(--shadow-medium);
    }

    .chat-message.user {
        background: var(--primary-gradient);
        color: white;
        margin-left: var(--spacing-xl);
    }

    .chat-message.assistant {
        margin-right: var(--spacing-xl);
        color: var(--text) !important;
    }

    .message-avatar {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: bold;
        font-size: 14px;
        margin-right: var(--spacing-md);
        flex-shrink: 0;
    }

    .user-avatar {
        background: rgba(255, 255, 255, 0.2);
        color: white;
    }

    .assistant-avatar {
        background: var(--primary-gradient);
        color: white;
    }

    .message-content {
        flex: 1;
        line-height: 1.6;
    }

    .message-time {
        font-size: 0.75rem;
        color: var(--text-muted);
        margin-top: var(--spacing-xs);
        opacity: 0.7;
    }

    .chat-message.user .message-time {
        color: rgba(255, 255, 255, 0.7);
    }

    /* Status indicator */
    .status-indicator {
        display: flex;
        align-items: center;
        gap: var(--spacing-sm);
        padding: var(--spacing-md);
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        border: 1px solid var(--glass-border);
        border-radius: var(--border-radius);
        box-shadow: var(--shadow-light);
        margin: var(--spacing-md) 0;
    }

    .status-dot {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background: var(--accent-color);
        animation: pulse 2s infinite;
    }

    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
    }

    /* Progress container */
    .progress-container {
        width: 100%;
        height: 4px;
        background: var(--bg-tertiary);
        border-radius: 2px;
        overflow: hidden;
        margin: var(--spacing-md) 0;
    }

    .progress-bar {
        height: 100%;
        background: var(--primary-gradient);
        border-radius: 2px;
        transition: width 0.3s ease;
    }

    /* Button styles */
    .stButton button {
        background: var(--primary-gradient) !important;
        color: white !important;
        border: none !important;
        border-radius: var(--border-radius) !important;
        padding: var(--spacing-sm) var(--spacing-md) !important;
        font-weight: 500 !important;
        transition: var(--transition) !important;
        box-shadow: var(--shadow-light) !important;
    }

    .stButton button:hover {
        transform: translateY(-2px) !important;
        box-shadow: var(--shadow-medium) !important;
    }

    .stButton button:active {
        transform: translateY(0) !important;
    }

    /* Secondary button */
    .stButton button[kind="secondary"] {
        background: var(--bg-tertiary) !important;
        color: var(--text) !important;
    }

    /* Delete chat button */
    .stButton button[title="Удалить чат"] {
        background: black !important;
        color: white !important;
    }

    /* Sidebar collapse button */
    button[data-testid="stSidebarCollapseButton"] {
        background: black !important;
        color: white !important;
        border: none !important;
        border-radius: 4px !important;
        padding: 4px 8px !important;
    }

    /* Form styles */
    .stTextInput input, .stTextArea textarea {
        background: var(--glass-bg) !important;
        border: 1px solid var(--glass-border) !important;
        border-radius: var(--border-radius) !important;
        color: var(--text) !important;
        transition: var(--transition) !important;
    }

    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.1) !important;
    }

    /* Scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
    }

    ::-webkit-scrollbar-track {
        background: var(--bg-tertiary);
        border-radius: 4px;
    }

    ::-webkit-scrollbar-thumb {
        background: var(--text-muted);
        border-radius: 4px;
    }

    ::-webkit-scrollbar-thumb:hover {
        background: var(--text-secondary);
    }

    /* Responsive design */
    @media (max-width: 768px) {
        .main-title {
            font-size: 2rem;
        }

        .subtitle {
            font-size: 1rem;
        }

        .chat-message {
            padding: var(--spacing-sm);
        }

        .message-avatar {
            width: 32px;
            height: 32px;
            font-size: 12px;
        }
    }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)


def init_theme():
    """Initialize theme settings"""
    if "theme" not in st.session_state:
        st.session_state.theme = "light"


def toggle_theme():
    """Toggle between light and dark theme"""
    if st.session_state.theme == "light":
        st.session_state.theme = "dark"
    else:
        st.session_state.theme = "light"
    apply_theme()


def apply_theme():
    """Apply the current theme to the page"""
    theme = st.session_state.get("theme", "light")
    if theme == "dark":
        bg_primary = "#0f172a"
        text = "#f8fafc"
        text_secondary = "#cbd5e1"
        bg_secondary = "#1e293b"
        bg_tertiary = "#334155"
        glass_bg = "rgba(15, 23, 42, 0.95)"
        glass_border = "rgba(255, 255, 255, 0.1)"
    else:
        bg_primary = "#ffffff"
        text = "#1e293b"
        text_secondary = "#64748b"
        bg_secondary = "#f8fafc"
        bg_tertiary = "#f1f5f9"
        glass_bg = "rgba(255, 255, 255, 0.95)"
        glass_border = "rgba(255, 255, 255, 0.2)"

    st.markdown(f"""
    <style>
    :root {{
        --bg-primary: {bg_primary};
        --text: {text};
        --text-secondary: {text_secondary};
        --bg-secondary: {bg_secondary};
        --bg-tertiary: {bg_tertiary};
        --glass-bg: {glass_bg};
        --glass-border: {glass_border};
    }}
    </style>
    """, unsafe_allow_html=True)


def load_chat_history():
    """Load chat history from JSON file"""
    if os.path.exists("chat_history.json"):
        try:
            with open("chat_history.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                # Ensure required keys exist
                if "chats" not in data:
                    data["chats"] = {}
                if "current_chat_id" not in data:
                    data["current_chat_id"] = None
                return data
        except Exception as e:
            st.error(f"Ошибка загрузки истории чатов: {e}")
    return {"chats": {}, "current_chat_id": None}


def save_chat_history(chat_data):
    """Save chat history to JSON file"""
    try:
        with open("chat_history.json", "w", encoding="utf-8") as f:
            json.dump(chat_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Ошибка сохранения истории чатов: {e}")


def create_new_chat():
    """Create a new chat session"""
    chat_id = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    chat = {
        "id": chat_id,
        "title": "Новый чат",
        "messages": [],
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    return chat_id, chat


def update_chat_title(chat_id, messages):
    """Update chat title based on first user message"""
    if not messages:
        return "Новый чат"

    # Find first user message
    for message in messages:
        if message["role"] == "user":
            content = message["content"].strip()
            # Create title from first 50 characters
            if len(content) > 50:
                title = content[:47] + "..."
            else:
                title = content
            return title

    return "Новый чат"


def check_word_export_request(prompt):
    """Check if the prompt contains a request to export to Word"""
    export_keywords = [
        "экспортируй в word", "экспорт в word", "word", "документ word",
        "скачать word", "word документ", "экспорт", "export to word",
        "word export", "word document"
    ]
    prompt_lower = prompt.lower().strip()
    return any(keyword in prompt_lower for keyword in export_keywords)


def generate_word_document(response, prompt, mode_name):
    """Generate a Word document from the response"""
    doc = Document()

    # Add title
    title = doc.add_heading('Ответ системы поиска и анализа нормативов', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    doc.add_paragraph(f'Режим: {mode_name}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    doc.add_paragraph('')

    # Add user prompt
    doc.add_heading('Запрос пользователя:', level=2)
    doc.add_paragraph(prompt)
    doc.add_paragraph('')

    # Add response
    doc.add_heading('Ответ системы:', level=2)
    doc.add_paragraph(response)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
